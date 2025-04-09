import re
from docx import Document
from datetime import datetime
import os
import argparse
import glob
import logging
import sys

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def parse_analysis_file(file_path):
    """
    Parses the content of an analysis results text file.

    Args:
        file_path (str): Path to the input .txt file.

    Returns:
        dict: A dictionary containing parsed 'header', 'metrics', and 'groups',
              or None if parsing fails significantly.
    """
    logging.info(f"Parsing analysis file: {file_path}")
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    except FileNotFoundError:
        logging.error(f"Input file not found: {file_path}")
        return None
    except Exception as e:
        logging.error(f"Error reading file {file_path}: {e}")
        return None

    parsed_data = {"header": {}, "metrics": {}, "groups": []}

    # --- Parse Header ---
    header_patterns = {
        'source_file': r"Source File:\s*(.*)",
        'model_used': r"Model Used:\s*(.*)",
        'analysis_date': r"Analysis Date:\s*(.*)"
    }
    for key, pattern in header_patterns.items():
        match = re.search(pattern, content)
        if match:
            parsed_data["header"][key] = match.group(1).strip()
        else:
             logging.warning(f"Could not find header key '{key}' in {file_path}")
             parsed_data["header"][key] = "N/A" # Default value

    # --- Parse Metrics ---
    metrics_match = re.search(r"ANALYSIS METRICS:\s*\n=+\s*\n(.*?)(?=\n={5,}\s*Connection Group|\Z)", content, re.DOTALL | re.MULTILINE) # Adjusted end pattern
    if metrics_match:
        metrics_block = metrics_match.group(1).strip()
        # Split into lines and parse key-value pairs more carefully
        for line in metrics_block.split('\n'):
            line = line.strip()
            if ':' in line:
                 # Split only on the *first* colon
                 key, value = line.split(':', 1)
                 # Clean up key (remove leading bullets/hyphens/numbers)
                 cleaned_key = re.sub(r"^\s*[-\*]?\s*\d*\.?\s*", "", key).strip()
                 # Further check: Ensure the cleaned key makes sense and isn't a sub-header itself
                 if cleaned_key and not cleaned_key.replace('.','',1).isdigit(): # Basic check to skip lines like "1." or "2."
                    parsed_data["metrics"][cleaned_key] = value.strip()
    else:
        logging.warning(f"Could not find 'ANALYSIS METRICS' block in {file_path}")


    # --- Parse Connection Groups ---
    analysis_content_start = metrics_match.end() if metrics_match else 0 # Start searching after metrics
    group_pattern = re.compile(
        r"={5,}\s*Connection Group\s*(\d+)\s*Analysis\s*={5,}\n"
        r"(.*?)"
        r"(?=\n={5,}\s*Connection Group|\Z)",
        re.DOTALL | re.MULTILINE
    )
    group_matches = group_pattern.finditer(content, pos=analysis_content_start)
    groups_found_count = 0
    for match in group_matches:
        groups_found_count += 1
        group_id = match.group(1)
        group_content = match.group(2).strip()
        if group_content:
            parsed_data["groups"].append({
                "id": group_id,
                "content": group_content
            })

    if groups_found_count == 0: # Check if the iterator yielded anything
         logging.warning(f"Could not find any 'Connection Group' blocks in {file_path}")

    logging.info(f"Parsed {len(parsed_data['groups'])} connection groups.")
    return parsed_data

def generate_docx_report(parsed_data, output_path):
    """
    Generates a .docx report from the parsed analysis data.

    Args:
        parsed_data (dict): The dictionary returned by parse_analysis_file.
        output_path (str): The full path to save the .docx file.
    """
    if not parsed_data:
        logging.error("Cannot generate report: Parsed data is missing.")
        return

    logging.info(f"Generating DOCX report: {output_path}")
    doc = Document()

    # --- Title ---
    doc.add_heading('Network Anomaly Connection Analysis Report', level=0)

    # --- Header / Summary Info ---
    doc.add_heading('Analysis Details', level=1)
    doc.add_paragraph(f"Source Analysis File: {parsed_data['header'].get('source_file', 'N/A')}")
    doc.add_paragraph(f"Analysis Generated On: {parsed_data['header'].get('analysis_date', 'N/A')}")
    doc.add_paragraph(f"Analysis Model Used: {parsed_data['header'].get('model_used', 'N/A')}")
    doc.add_paragraph(f"Total Connection Groups Analyzed in Source: {len(parsed_data['groups'])}")
    doc.add_paragraph() # Add a blank line

    # --- Metrics Section ---
    if parsed_data['metrics']:
        doc.add_heading('Analysis Metrics Summary', level=1)
        # Sort metrics for consistent order (optional)
        for key, value in sorted(parsed_data["metrics"].items()):
            # Add bold formatting to the key (optional)
            p = doc.add_paragraph()
            p.add_run(f"{key}:").bold = True
            p.add_run(f" {value}")
        doc.add_paragraph() # Add a blank line
    else:
        logging.warning("No metrics data found to include in the report.")


    # --- Connection Groups Section ---
    if parsed_data['groups']:
        doc.add_heading('Detailed Connection Group Analysis', level=1)
        for group in parsed_data["groups"]:
            doc.add_heading(f"Connection Group {group['id']} Analysis", level=2)

            # --- Clean the content before adding it ---
            cleaned_content = []
            for line in group['content'].split('\n'):
                # Remove lines consisting only of '#' characters (allowing whitespace)
                if not re.match(r"^\s*#+\s*$", line):
                    cleaned_content.append(line)
            content_to_add = "\n".join(cleaned_content)
            # --- End Cleaning ---

            # Add the cleaned content from the LLM for this group
            doc.add_paragraph(content_to_add) # Use cleaned content
            doc.add_paragraph() # Add a blank line between groups
    else:
        logging.warning("No connection group analysis found to include in the report.")


    # --- Save the report ---
    try:
        output_dir = os.path.dirname(output_path)
        os.makedirs(output_dir, exist_ok=True)
        doc.save(output_path)
        logging.info(f"DOCX Report saved successfully to: {output_path}")
    except Exception as e:
        logging.error(f"Error saving DOCX report to {output_path}: {e}")


if __name__ == "__main__":
    # --- Argument Parsing ---
    script_dir = os.path.dirname(os.path.abspath(__file__))
    default_analysis_dir = os.path.join(script_dir, "analysis_results")
    default_docx_dir = os.path.join(script_dir, "docx_reports")

    parser = argparse.ArgumentParser(description="Generate DOCX reports from Ollama analysis text files.")
    parser.add_argument("-i", "--input-dir", default=default_analysis_dir,
                        help="Directory containing the analysis text files.")
    parser.add_argument("-o", "--output-dir", default=default_docx_dir,
                        help="Directory to save the generated DOCX reports.")
    parser.add_argument("-f", "--file", default=None,
                        help="Process only a specific analysis text file (full path or relative to input-dir).")
    parser.add_argument("--mode", choices=["latest", "all", "file"], default="latest",
                        help="Processing mode: 'latest' analysis file, 'all' files in input-dir, or specific 'file'.")

    args = parser.parse_args()

    # --- Input File Selection ---
    files_to_process = []
    input_dir = args.input_dir

    if not os.path.isdir(input_dir):
        logging.error(f"Input directory not found: {input_dir}")
        sys.exit(1)

    mode = args.mode
    if args.file:
        mode = "file"

    if mode == "file":
        if args.file:
            specific_file_path = args.file
            if not os.path.isabs(specific_file_path):
                specific_file_path = os.path.join(input_dir, specific_file_path)
            if os.path.exists(specific_file_path) and specific_file_path.endswith('.txt'):
                files_to_process.append(specific_file_path)
            else:
                logging.error(f"Specified analysis file not found or invalid: {specific_file_path}")
        else:
            logging.error("Mode 'file' selected, but no file specified with -f argument.")

    elif mode == "latest":
        analysis_files = glob.glob(os.path.join(input_dir, "*_analysis_*.txt"))
        if analysis_files:
            try:
                latest_file = max(analysis_files, key=os.path.getmtime)
                files_to_process.append(latest_file)
                logging.info(f"Processing latest analysis file: {os.path.basename(latest_file)}")
            except ValueError: # Handle case where glob returns empty list after check (race condition?)
                 logging.warning(f"No analysis files (*_analysis_*.txt) found in {input_dir}")
        else:
            logging.warning(f"No analysis files (*_analysis_*.txt) found in {input_dir}")

    elif mode == "all":
        analysis_files = glob.glob(os.path.join(input_dir, "*_analysis_*.txt"))
        if analysis_files:
            files_to_process.extend(sorted(analysis_files)) # Process in sorted order
            logging.info(f"Processing all {len(files_to_process)} analysis files found in {input_dir}")
        else:
            logging.warning(f"No analysis files (*_analysis_*.txt) found in {input_dir}")


    # --- Processing Loop ---
    output_dir = args.output_dir
    os.makedirs(output_dir, exist_ok=True)

    if not files_to_process:
        logging.info("No analysis files selected for processing.")
    else:
        processed_count = 0
        for input_file_path in files_to_process:
            logging.info(f"--- Processing: {os.path.basename(input_file_path)} ---")
            parsed_data = parse_analysis_file(input_file_path)

            if parsed_data and (parsed_data['metrics'] or parsed_data['groups']): # Check if any useful data was parsed
                # Generate output filename
                base_input_name = os.path.splitext(os.path.basename(input_file_path))[0]
                report_timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"{base_input_name}_Report_{report_timestamp}.docx"
                output_file_path = os.path.join(output_dir, output_filename)

                generate_docx_report(parsed_data, output_file_path)
                processed_count += 1
            elif parsed_data:
                 logging.warning(f"Skipping DOCX generation for {os.path.basename(input_file_path)} as no metrics or groups were successfully parsed.")
            else:
                logging.error(f"Skipping DOCX generation for {os.path.basename(input_file_path)} due to parsing errors.")
            logging.info("-" * (len(os.path.basename(input_file_path)) + 18)) # Separator

    logging.info(f"DOCX report generation process finished. {processed_count} reports generated.")
