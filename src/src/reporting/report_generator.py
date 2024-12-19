import re
from docx import Document
from datetime import datetime
import os

# Get the base directory of the project (go up two levels from the script)
base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))

# Construct the absolute path to the input file
input_file = os.path.join(base_dir, 'data', 'streamed_responses', 'granite3-dense:2b_combined_output_20241217_1955.txt')

# Ensure the input file exists
if not os.path.exists(input_file):
    raise FileNotFoundError(f"Input file not found at: {input_file}")

# Construct the directory for the reports
reports_dir = os.path.join(base_dir, 'data', 'reports')
os.makedirs(reports_dir, exist_ok=True)

# Generate the filename with the date and time
timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
output_file = os.path.join(reports_dir, f"ICS_Report_{timestamp}.docx")

def generate_report(parsed_data, output_path):
    doc = Document()
    doc.add_heading('ICS Anomaly Connection Report', 0)

    # Summary section
    doc.add_heading('Summary', level=1)
    summary = (
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Total Conversations Analyzed: {len(parsed_data['conversations'])}"
    )
    doc.add_paragraph(summary)

    # Metrics section
    doc.add_heading('Consolidated Metrics', level=1)
    for key, value in parsed_data["metrics"].items():
        doc.add_paragraph(f"{key}: {value}")

    # Conversations section
    doc.add_heading('Connection Conversations', level=1)
    for convo in parsed_data["conversations"]:
        doc.add_heading(f"Conversation {convo['id']}", level=2)
        doc.add_paragraph(convo['details'])

    # Save the report
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

def parse_output(file_path):
    with open(file_path, 'r') as file:
        lines = file.read()

    data = {"metrics": {}, "conversations": []}

    # Parse consolidated metrics
    metrics_pattern = r"Consolidated Metrics:\n=+\n(.*?)(?=\nConversation|\Z)"
    metrics_block = re.search(metrics_pattern, lines, re.DOTALL)
    if metrics_block:
        metrics_lines = metrics_block.group(1).split("\n")
        for line in metrics_lines:
            if ":" in line:
                key, value = map(str.strip, line.split(":", 1))
                data["metrics"][key] = value

    # Parse conversations
    conversation_pattern = r"Conversation \d+:\n(.*?)(?=(Conversation \d+:|$))"
    conversations = re.findall(conversation_pattern, lines, re.DOTALL)
    
    for idx, convo in enumerate(conversations):
        details = convo[0].strip().split("\n")
        conversation_info = {"id": idx + 1, "details": " ".join(details)}
        data["conversations"].append(conversation_info)

    return data

# Process
parsed_data = parse_output(input_file)
generate_report(parsed_data, output_file)
